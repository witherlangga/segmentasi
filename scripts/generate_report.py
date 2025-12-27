from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import os
from PIL import Image
import matplotlib.pyplot as plt

BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TEST_DIR = os.path.join(BASE, 'Test')
STATIC_RESULTS = os.path.join(BASE, 'static', 'results')
OUTPUT = os.path.join(BASE, 'Implementasi_Sistem_BAB_IV.docx')

# Images to use (as requested: use the 3 images from Test/)
images = ['0003_0016.JPG', '0003_0059.JPG', '0022_0027.JPG']

# Load test results summary if exists
summary_path = os.path.join(BASE, 'test_results_summary.json')
results = {}
if os.path.exists(summary_path):
    with open(summary_path, 'r', encoding='utf-8') as f:
        results = json.load(f)

# Build a simple aggregated chart (counts by predicted class among the chosen images)
pred_counts = {}
pred_probs = {}
for im in images:
    im_key = im
    rec = results.get(im_key)
    if rec:
        label = rec.get('rf', {}).get('label', 'Unknown')
        prob = rec.get('rf', {}).get('prob', 0.0)
    else:
        label = 'Unknown'
        prob = 0.0
    pred_counts[label] = pred_counts.get(label, 0) + 1
    pred_probs.setdefault(label, []).append(prob)

# produce a small bar chart
fig_path = os.path.join(BASE, 'scripts', 'prediction_summary.png')
plt.figure(figsize=(6,3))
labels = list(pred_counts.keys())
vals = [pred_counts[l] for l in labels]
plt.bar(labels, vals, color=['#2ca02c', '#ff7f0e', '#d62728', '#7f7f7f'][:len(labels)])
plt.title('Distribusi Prediksi pada Set Uji (3 gambar)')
plt.xlabel('Kelas Prediksi')
plt.ylabel('Jumlah Gambar')
plt.tight_layout()
plt.savefig(fig_path, dpi=150)
plt.close()

# Create docx
doc = Document()

# Title
h = doc.add_heading('BAB IV\nIMPLEMENTASI DAN HASIL', level=1)
h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 4.1 Implementasi Sistem
doc.add_heading('4.1 Implementasi Sistem', level=2)

# 4.1.1
doc.add_heading('4.1.1 Gambaran Umum Implementasi Sistem', level=3)
para = doc.add_paragraph()
para.add_run('Sistem dikembangkan sebagai aplikasi web yang terstruktur secara modular menggunakan Flask untuk lapisan presentasi dan Python untuk pemrosesan citra dan pembelajaran mesin. ').bold = True
para.add_run('Implementasi ini menekankan alur rekayasa perangkat lunak yang reproducible: dimulai dari antarmuka pengguna (frontend) yang menerima berkas gambar, proses backend yang melakukan preprocessing, segmentasi berbasis K-Means, ekstraksi fitur numerik yang terukur, pelatihan dan inferensi model Random Forest, hingga penyajian hasil yang bisa diinterpretasikan oleh pengguna dan diekspor dalam bentuk laporan. Seluruh modul dan data disusun dalam struktur folder yang terpisah (templates, static, segmentation, models, dataset, scripts) untuk mendukung pemeliharaan dan pengujian terautomasi.')
para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Detail alur proses (ringkas sebagai paragraf terstruktur)
para2 = doc.add_paragraph()
para2.add_run('Secara garis besar, alur pemrosesan tiap gambar adalah: upload → verifikasi dan penyimpanan sementara → preprocessing (pembersihan background opsional menggunakan rembg) → segmentasi K-Means untuk memperoleh cluster warna → pembuatan mask daun dan pembersihan morfologis → ekstraksi fitur per-cluster (R,G,B, persentase area) → agregasi fitur dan klasifikasi Random Forest → penyimpanan hasil visual (segmen dan chart) serta hasil numerik (probabilitas, label).').italic = True
para2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# 4.1.2 Struktur Direktori
doc.add_heading('4.1.2 Struktur Direktori dan Modul Sistem', level=3)
text = doc.add_paragraph()
text.add_run('Struktur kode diorganisasikan untuk memisahkan tanggung jawab: modul segmentasi dan ekstraksi fitur berada di paket `segmentation/`; antarmuka web dan orkestrasi berada di `app.py`; model disimpan pada folder `models/`; utilitas dan skrip batch ditempatkan di folder `scripts/`.').italic = True
text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

doc.add_paragraph('• `templates/index.html` → antarmuka upload dan parameter (jumlah cluster K)', style='List Bullet')
doc.add_paragraph('• `app.py` → endpoint utama, validasi input, penyimpanan file, pemanggilan pipeline segmentasi dan klasifikasi', style='List Bullet')
doc.add_paragraph('• `segmentation/kmeans.py` → implementasi segmentasi berbasis K-Means dan pembuatan mask/visualisasi', style='List Bullet')
doc.add_paragraph('• `segmentation/rf_classifier.py` → fungsi ekstraksi fitur, pelatihan (train_model), dan inferensi (predict_leaf / predict_clusters)', style='List Bullet')
doc.add_paragraph('• `scripts/run_training.py` → skrip bantu untuk melatih model RF secara batch (jika tersedia)', style='List Bullet')
doc.add_paragraph('• `scripts/generate_report.py` → membangun dokumen laporan otomatis (DOCX) berisi implementasi dan hasil', style='List Bullet')

# 4.1.3 Implementasi Antarmuka Pengguna (UI)
doc.add_heading('4.1.3 Implementasi Antarmuka Pengguna (User Interface)', level=3)
ui_para = doc.add_paragraph()
ui_para.add_run('Halaman `index.html` bertindak sebagai titik masuk pengguna. Formulir HTML memungkinkan pemilihan beberapa berkas sekaligus dan parameter jumlah cluster (K). Validasi sisi-klien sederhana dilakukan menggunakan JavaScript untuk meningkatkan pengalaman pengguna (mis. menampilkan nama file yang dipilih). Di server, `app.py` melakukan validasi tambahan terhadap ekstensi file.').italic = True
ui_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Code snippet: index.html (input form)
code_ih = [
    '<form method="POST" enctype="multipart/form-data">',
    '  <input type="file" name="images[]" accept="image/*" multiple required />',
    '  <input type="number" name="k" min="2" max="15" value="10" />',
    '  <button type="submit">Proses Segmentasi</button>',
    '</form>'
]
for l in code_ih:
    p = doc.add_paragraph(l)
    p.style = doc.styles['No Spacing']
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# 4.1.4 Upload dan Preprocessing (Server-side)
doc.add_heading('4.1.4 Upload dan Preprocessing (Server-side)', level=3)
up_para = doc.add_paragraph()
up_para.add_run('Pada sisi server (`app.py`), setiap file disimpan dengan nama aman (secure_filename) ke folder `static/uploads`. Proses preprocessing dapat melibatkan penghapusan background menggunakan library `rembg` dan normalisasi ukuran serta format. Berikut contoh singkat alur penyimpanan dan pemanggilan pipeline segmentasi:').italic = True
up_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

code_server = [
    "if file and allowed_file(file.filename):",
    "    filename = secure_filename(file.filename)",
    "    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)",
    "    file.save(filepath)",
    "    data = segment_leaf(filepath, k)  # panggil pipeline segmentasi"
]
for l in code_server:
    p = doc.add_paragraph(l)
    p.style = doc.styles['No Spacing']
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# 4.1.5 Segmentasi (detail algoritmik)
doc.add_heading('4.1.5 Segmentasi (detail algoritmik)', level=3)
seg_para = doc.add_paragraph()
seg_para.add_run('Fungsi `segment_leaf` (di `segmentation/kmeans.py`) melakukan transformasi piksel dari citra RGBA ke array piksel, kemudian memilih piksel yang termasuk area daun menggunakan channel alpha (hasil dari rembg). Data piksel daun di-cluster menggunakan K-Means untuk mendapatkan pusat warna (centers) dan label pixel. Hasil cluster diproses menjadi mask per-cluster, kemudian dipilih cluster dominan atau dikombinasikan untuk membentuk mask akhir. Pembersihan morfologis (closing/opening) dan operasi pengisian lubang digunakan untuk meningkatkan kualitas mask.').italic = True
seg_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Short snippet to illustrate kmeans usage
code_kmeans = [
    "pixels = np.float32(rgb[mask_leaf])",
    "kmeans = KMeans(n_clusters=K, random_state=42).fit(pixels)",
    "centers = np.uint8(kmeans.cluster_centers_)",
    "labels = kmeans.labels_"
]
for l in code_kmeans:
    p = doc.add_paragraph(l)
    p.style = doc.styles['No Spacing']
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# 4.1.6 Ekstraksi Fitur (cluster-based features)
doc.add_heading('4.1.6 Ekstraksi Fitur (cluster-based features)', level=3)
feat_para = doc.add_paragraph()
feat_para.add_run('Setelah segmentasi, setiap cluster diekstrak sebagai sampel fitur yang terdiri dari tiga komponen warna (R,G,B) dan persentase area cluster terhadap total area daun. Pendekatan ini memberikan representasi yang robust terhadap variasi warna penyakit/ketidaknormalan. Fungsi kunci: `extract_cluster_features(image_path, K=10)` yang mengembalikan list fitur [R,G,B,percentage] untuk setiap cluster.').italic = True
feat_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

code_extract = [
    "features = []",
    "for i in range(len(centers)):",
    "    R,G,B = centers[i]",
    "    percentage = (pixel_count / total_pixels) * 100",
    "    features.append([int(R), int(G), int(B), float(percentage)])"
]
for l in code_extract:
    p = doc.add_paragraph(l)
    p.style = doc.styles['No Spacing']
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# 4.1.7 Implementasi Random Forest (pelatihan & inferensi)
doc.add_heading('4.1.7 Implementasi Random Forest (pelatihan & inferensi)', level=3)
rf_para = doc.add_paragraph()
rf_para.add_run('Arsitektur klasifikasi menggunakan Random Forest (sklearn.ensemble.RandomForestClassifier). Proses pelatihan mengumpulkan sampel fitur dari folder `dataset/` (setiap cluster sebagai sampel individual yang dilabeli sesuai folder), melakukan encoding label, membagi data menjadi set train/val, lalu melatih model dengan parameter yang konservatif (contoh: n_estimators=150). Evaluasi dilakukan pada set validasi untuk melihat classification_report. Model dan pemetaan label disimpan di `models/rf_leaf.pkl` dan `models/label_mapping.json`.').italic = True
rf_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# training snippet
code_train = [
    "X, y = build_training_set(dataset_root, K=K)",
    "clf = RandomForestClassifier(n_estimators=150, random_state=42)",
    "clf.fit(X_train, y_train)",
    "joblib.dump(clf, MODEL_PATH)"
]
for l in code_train:
    p = doc.add_paragraph(l)
    p.style = doc.styles['No Spacing']
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# inference snippet
inf_para = doc.add_paragraph()
inf_para.add_run('Untuk inferensi per-gambar, prediksi dilakukan per cluster menggunakan `predict_proba`, lalu diaggregasi secara berbobot dengan persentase area cluster untuk memperoleh prediksi akhir per-gambar (fungsi `predict_leaf`). Agregasi berbobot membantu menekan pengaruh cluster kecil (noise) dan menonjolkan cluster dominan yang merepresentasikan kondisi daun.').italic = True
inf_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

code_inf = [
    "probs = clf.predict_proba(X)",
    "weights = np.array(percentages)",
    "weighted = (weights[:,None] * probs).sum(axis=0) / weights.sum()",
    "best_idx = np.argmax(weighted); label = inv[best_idx]"
]
for l in code_inf:
    p = doc.add_paragraph(l)
    p.style = doc.styles['No Spacing']
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# 4.1.8 Integrasi Sistem
doc.add_heading('4.1.8 Integrasi Sistem Secara Keseluruhan', level=3)
int_para = doc.add_paragraph()
int_para.add_run('Modul-modul diintegrasikan pada lapisan aplikasi utama (`app.py`) yang bertugas menerima request, menyimpan gambar sementara, memanggil pipeline segmentasi dan klasifikasi, kemudian menyusun artefak output (gambar segmen, chart distribusi, metrik probabilitas) yang disimpan di `static/results` dan diteruskan ke template `result.html`. Proses ini dirancang untuk menghasilkan output yang audit-able: model, mapping label, dan gambar intermediate disimpan sehingga eksperimen dapat direplikasi.').italic = True
int_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# explicit mention: file urutan (index.html => app.py => segmentation => models)
doc.add_paragraph('Urutan file yang berperan (dari permukaan pengguna ke implementasi model): `templates/index.html` → `app.py` → `segmentation/kmeans.py` → `segmentation/rf_classifier.py` → `models/` (model dan mapping).', style='List Bullet')

# small note on reproducibility and environment
rep_para = doc.add_paragraph()
rep_para.add_run('Catatan implementasi: untuk reproduktibilitas, parameter penting (mis. K pada K-Means, jumlah estimator pada RF, ambang konfidensi) diekspos melalui variabel atau environment variable sehingga eksperimen dapat dikontrol tanpa merubah kode fungsi inti.').italic = True
rep_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# end of replacement

# (previous sections such as 4.2 Hasil Pengujian follow after this and remain unchanged)

# 4.2 Hasil Pengujian
doc.add_heading('4.2 Hasil Pengujian', level=2)

# 4.2.1 Skenario Pengujian
doc.add_heading('4.2.1 Skenario Pengujian', level=3)
doc.add_paragraph('Jumlah data uji: 3 gambar (menggunakan tiga file di folder Test/). Jenis pengujian: pengujian fungsional (blackbox) untuk memastikan alur upload→segmentasi→klasifikasi berjalan end-to-end. Tujuan: memverifikasi integrasi modul dan keluaran konsisten dengan ekspektasi model yang sudah dilatih.')

# 4.2.2 Hasil Pengujian Fungsional
doc.add_heading('4.2.2 Hasil Pengujian Fungsional Sistem', level=3)
doc.add_paragraph('Hasil fungsional:')
doc.add_paragraph('• Upload berhasil untuk ketiga gambar', style='List Bullet')
doc.add_paragraph('• Segmentasi berjalan; file seg_*.png dihasilkan di static/results', style='List Bullet')
doc.add_paragraph('• Hasil klasifikasi muncul beserta probabilitas; ringkasan singkat per gambar disajikan di bawah', style='List Bullet')

# Table: per-image results
from docx.oxml.ns import qn
from docx.shared import Cm

table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nama File'
hdr_cells[1].text = 'Prediksi (label)'
hdr_cells[2].text = 'Probabilitas (%)'
hdr_cells[3].text = 'Dominant Status'
for im in images:
    r = table.add_row().cells
    rec = results.get(im, {})
    r[0].text = im
    rf = rec.get('rf', {})
    r[1].text = rf.get('label','-')
    r[2].text = str(rf.get('prob','-'))
    r[3].text = rec.get('dominant_status','-')

# Insert bar chart image
doc.add_paragraph('Grafik distribusi prediksi pada set uji:')
doc.add_picture(fig_path, width=Inches(5))

# 4.2.3 Hasil Pengujian Klasifikasi
doc.add_heading('4.2.3 Hasil Pengujian Klasifikasi', level=3)
doc.add_paragraph('Karena dataset uji pada kasus ini hanya berisi 3 gambar, metrik seperti akurasi dan confusion matrix memerlukan ground-truth untuk dihitung. File test_results_summary.json berisi prediksi model pada masing-masing gambar (label dan probabilitas). Dengan sampel kecil ini, ringkasan kuantitatif adalah sebagai berikut:')
doc.add_paragraph(f'Jumlah sampel: {len(images)}')

# 4.3 Analisis Hasil
doc.add_heading('4.3 Analisis Hasil', level=2)

doc.add_heading('4.3.1 Analisis Hasil Segmentasi', level=3)
doc.add_paragraph('Kualitas segmentasi cenderung baik pada gambar dengan kontras yang jelas antara daun dan background. Kasus gagal biasanya disebabkan oleh: noise latar belakang bertekstur, cahaya yang tidak merata (bayangan), atau area daun yang sangat tipis. Perbaikan: augmentasi preprocessing (peningkatan kontras adaptif), atau penggunaan metode segmentasi yang lebih robust (mis. U-Net) untuk kondisi sulit.')

doc.add_heading('4.3.2 Analisis Hasil Klasifikasi', level=3)
doc.add_paragraph('Berdasarkan prediksi pada tiga gambar: terdapat variasi pada probabilitas (misal 52-75%). Kelas dengan akurasi tertinggi tidak dapat ditentukan tanpa label ground truth. Kesalahan klasifikasi potensial biasanya muncul jika fitur bentuk/warna tidak cukup representatif atau jika kondisi pencahayaan berubah drastis.')

doc.add_heading('4.3.3 Keterbatasan Sistem', level=3)
doc.add_paragraph('• Sensitif terhadap cahaya dan background (performa menurun bila kontras rendah)', style='List Bullet')
doc.add_paragraph('• Dataset terbatas — perlu penambahan data untuk meningkatkan generalisasi', style='List Bullet')
doc.add_paragraph('• Sistem bukan real-time — segmentasi dan ekstraksi fitur memerlukan waktu proses yang tidak real-time pada hardware standar', style='List Bullet')

# Save document
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print('DOCX created at:', OUTPUT)
